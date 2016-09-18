__author__ = 'schrecknetuser'

import os
import shutil
from docx import Document
import locale


    #document = Document(file)
    #print document.paragraphs[0].text

def convert_and_create_contents(dir):

    print(dir)

    renamed_folder = dir + '_rename'
    if os.path.exists(renamed_folder):
        shutil.rmtree(renamed_folder)
    os.makedirs(renamed_folder)

    i = 1

    for file in sorted(os.listdir(dir), key=str.lower):
        if not file.endswith('docx'):
            continue
        shutil.copyfile(os.path.join(dir, file), os.path.join(renamed_folder, str(i).zfill(4) + '.docx'))
        i += 1

    with open(os.path.join(renamed_folder, 'config.yml'), 'w') as f:
        f.write('pdf_regex: \n')
        f.write('html_regex: \'^(.+).html$\'\n')
        f.write('book:\n')
        f.write('  has_pdf: \'false\'\n')
        f.write('  picture_path: \n')
        f.write('  book_file: \n')
        f.write('  name_prefix: \n')
        f.write('  tree_prefix: \n')
        f.write('  name: \n')
        f.write('  name_comment: \n')
        f.write('  page_count: \'%d\'\n' % (i - 1))
        f.write('  first_page: \'1\'\n')
        f.write('  synopsis: \n')
        f.write('  contents: \n')

        for file in sorted(os.listdir(renamed_folder), key=str.lower):
            if not file.endswith('docx'):
                continue
            document = Document(os.path.join(renamed_folder, file))
            string = ''
            page_number = str(int(file.split('.')[0]))
            j = 0
            for paragraph in document.paragraphs:
                if not paragraph.text:
                    continue
                #j += 1
                #if j == 1:
                #    continue
                string += "'%s'" % paragraph.text
                break
                #if j == 3:
                #    break
            f.write('  ' * 2 + '- \n')
            f.write('  ' * 3 + 'name: ' + string.replace("\'\'", " ") + '\n')
            f.write('  ' * 3 + 'page: ' + '\'' + page_number + '\'\n')
            f.write('  ' * 3 + 'type: ' + '\'page\'' + '\n')

def __main__():
    convert_and_create_contents('/users/schrecknetuser/Downloads/living_ethics_1')



if __name__ == '__main__':
    __main__()
